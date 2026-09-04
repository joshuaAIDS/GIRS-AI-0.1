"""
Document Intelligence & Local RAG Engine for IGIRS AI.
Parses, indexes, chunks, and answers queries across local PDFs, lecture notes, resumes, and code files.
"""
import io
import os
import re
import json
import math
import uuid
import base64
import logging
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional, Union, Tuple

import config

logger = logging.getLogger("IGIRS.DocumentEngine")

# Supported file extensions
EXT_CODE = {
    ".py": "Python", ".js": "JavaScript", ".ts": "TypeScript",
    ".html": "HTML", ".css": "CSS", ".java": "Java", ".cpp": "C++",
    ".c": "C", ".rs": "Rust", ".go": "Go", ".sql": "SQL",
    ".sh": "Shell Script", ".json": "JSON", ".yaml": "YAML",
    ".yml": "YAML", ".xml": "XML", ".php": "PHP", ".rb": "Ruby"
}
EXT_DOCS = {".pdf", ".docx", ".txt", ".md", ".rtf", ".csv", ".log"}

class DocumentEngine:
    def __init__(
        self,
        store_file: Path = config.KNOWLEDGE_STORE_FILE,
        docs_dir: Path = config.DOCUMENTS_DIR
    ):
        self.store_file = Path(store_file)
        self.docs_dir = Path(docs_dir)
        self.docs_dir.mkdir(parents=True, exist_ok=True)
        self.documents: Dict[str, Dict[str, Any]] = {}
        self._load_store()

    def _load_store(self):
        """Loads indexed documents metadata from disk."""
        if self.store_file.exists():
            try:
                with open(self.store_file, "r", encoding="utf-8") as f:
                    self.documents = json.load(f)
                logger.info(f"Loaded {len(self.documents)} indexed documents from knowledge store.")
            except Exception as e:
                logger.error(f"Failed to load knowledge store: {e}")
                self.documents = {}
        else:
            self.documents = {}

    def _save_store(self):
        """Persists indexed documents metadata to disk."""
        try:
            with open(self.store_file, "w", encoding="utf-8") as f:
                json.dump(self.documents, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Failed to persist knowledge store: {e}")

    # --- Extraction Helpers ---

    def _extract_pdf(self, file_path_or_bytes: Union[str, Path, bytes]) -> List[Dict[str, Any]]:
        """Extracts text per page from a PDF file."""
        pages = []
        try:
            import pypdf
            if isinstance(file_path_or_bytes, bytes):
                reader = pypdf.PdfReader(io.BytesIO(file_path_or_bytes))
            else:
                reader = pypdf.PdfReader(str(file_path_or_bytes))

            for idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                clean_text = text.strip()
                if clean_text:
                    pages.append({
                        "page_number": idx + 1,
                        "text": clean_text
                    })
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
        return pages

    def _extract_docx(self, file_path_or_bytes: Union[str, Path, bytes]) -> List[Dict[str, Any]]:
        """Extracts paragraphs and tables from a Word (.docx) file."""
        sections = []
        try:
            import docx
            if isinstance(file_path_or_bytes, bytes):
                doc = docx.Document(io.BytesIO(file_path_or_bytes))
            else:
                doc = docx.Document(str(file_path_or_bytes))

            full_text = []
            for p in doc.paragraphs:
                if p.text.strip():
                    full_text.append(p.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        full_text.append(" | ".join(row_cells))

            text = "\n\n".join(full_text)
            if text:
                sections.append({"page_number": 1, "text": text})
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
        return sections

    def _extract_code(self, file_path_or_bytes: Union[str, Path, bytes], ext: str) -> List[Dict[str, Any]]:
        """Extracts source code with line tracking and structure metadata."""
        if isinstance(file_path_or_bytes, bytes):
            try:
                content = file_path_or_bytes.decode("utf-8")
            except UnicodeDecodeError:
                content = file_path_or_bytes.decode("latin-1", errors="replace")
        else:
            try:
                with open(file_path_or_bytes, "r", encoding="utf-8") as f:
                    content = f.read()
            except UnicodeDecodeError:
                with open(file_path_or_bytes, "r", encoding="latin-1", errors="replace") as f:
                    content = f.read()

        lang = EXT_CODE.get(ext.lower(), "Code")
        lines = content.splitlines()
        chunks = []
        step = 100
        for i in range(0, max(1, len(lines)), step):
            sub_lines = lines[i:i + step]
            chunk_text = f"// Language: {lang} | Lines {i+1} to {min(len(lines), i+step)}\n" + "\n".join(sub_lines)
            chunks.append({
                "page_number": f"L{i+1}-{min(len(lines), i+step)}",
                "text": chunk_text
            })
        return chunks

    def _extract_text(self, file_path_or_bytes: Union[str, Path, bytes]) -> List[Dict[str, Any]]:
        """Extracts raw text / markdown."""
        if isinstance(file_path_or_bytes, bytes):
            try:
                content = file_path_or_bytes.decode("utf-8")
            except UnicodeDecodeError:
                content = file_path_or_bytes.decode("latin-1", errors="replace")
        else:
            try:
                with open(file_path_or_bytes, "r", encoding="utf-8") as f:
                    content = f.read()
            except UnicodeDecodeError:
                with open(file_path_or_bytes, "r", encoding="latin-1", errors="replace") as f:
                    content = f.read()

        return [{"page_number": 1, "text": content.strip()}]

    # --- Classification & Chunking ---

    def _classify_document(self, filename: str, full_text: str) -> str:
        """Determines the document category for tailored RAG system prompting."""
        fn_lower = filename.lower()
        txt_lower = full_text[:4000].lower()
        ext = Path(filename).suffix.lower()

        if ext in EXT_CODE:
            return "code"

        # Check for Resume / CV
        resume_keywords = ["education", "experience", "skills", "projects", "certifications", "linkedin", "curriculum vitae", "resume", "work experience", "gpa"]
        resume_score = sum(1 for kw in resume_keywords if kw in txt_lower or kw in fn_lower)
        if resume_score >= 3 or "resume" in fn_lower or "cv" in fn_lower:
            return "resume"

        # Check for Lecture Notes / Coursework
        lecture_keywords = ["lecture", "chapter", "syllabus", "slide", "homework", "exam", "course", "assignment", "instructor", "professor", "notes", "module", "definition"]
        lecture_score = sum(1 for kw in lecture_keywords if kw in txt_lower or kw in fn_lower)
        if lecture_score >= 2 or "lecture" in fn_lower or "notes" in fn_lower:
            return "lecture_notes"

        # Check for Research Paper
        paper_keywords = ["abstract", "introduction", "methodology", "references", "ieee", "arxiv", "conclusion", "dataset"]
        if sum(1 for kw in paper_keywords if kw in txt_lower) >= 3:
            return "research_paper"

        if ext == ".pdf":
            return "pdf_document"

        return "document"

    def _chunk_text(self, pages: List[Dict[str, Any]], chunk_size: int = 500, overlap: int = 60) -> List[Dict[str, Any]]:
        """Splits multi-page text into search chunks with source references."""
        all_chunks = []
        chunk_idx = 0

        for page in pages:
            page_ref = page.get("page_number", 1)
            text = page.get("text", "")
            words = text.split()

            if len(words) <= chunk_size:
                chunk_idx += 1
                all_chunks.append({
                    "id": chunk_idx,
                    "page": page_ref,
                    "text": text,
                    "word_count": len(words)
                })
            else:
                for i in range(0, len(words), chunk_size - overlap):
                    sub_words = words[i:i + chunk_size]
                    if not sub_words:
                        continue
                    chunk_idx += 1
                    all_chunks.append({
                        "id": chunk_idx,
                        "page": page_ref,
                        "text": " ".join(sub_words),
                        "word_count": len(sub_words)
                    })

        return all_chunks

    # --- Ingestion ---

    def ingest_file(
        self,
        source: Union[str, Path, bytes],
        filename: str,
        save_copy: bool = True
    ) -> Dict[str, Any]:
        """Ingests and indexes a local file or raw bytes into the Knowledge Vault."""
        p = Path(filename)
        ext = p.suffix.lower()
        doc_id = str(uuid.uuid4())[:8]

        pages = []
        if ext == ".pdf":
            pages = self._extract_pdf(source)
        elif ext == ".docx":
            pages = self._extract_docx(source)
        elif ext in EXT_CODE:
            pages = self._extract_code(source, ext)
        elif ext in EXT_DOCS or ext in [".txt", ".md", ".log", ".csv", ".json"]:
            pages = self._extract_text(source)
        else:
            pages = self._extract_text(source)

        if not pages:
            return {"status": "error", "message": f"Could not extract readable text from '{filename}'."}

        full_text = "\n\n".join(p.get("text", "") for p in pages)
        category = self._classify_document(filename, full_text)
        chunks = self._chunk_text(pages)

        stored_path = ""
        if save_copy:
            try:
                dest = self.docs_dir / f"{doc_id}_{p.name}"
                if isinstance(source, bytes):
                    with open(dest, "wb") as f:
                        f.write(source)
                elif Path(source).exists():
                    import shutil
                    shutil.copy2(str(source), str(dest))
                stored_path = str(dest)
            except Exception as e:
                logger.debug(f"Could not persist document copy: {e}")

        doc_record = {
            "id": doc_id,
            "filename": p.name,
            "category": category,
            "extension": ext,
            "stored_path": stored_path,
            "total_pages": len(pages),
            "total_chunks": len(chunks),
            "total_words": len(full_text.split()),
            "created_at": datetime.now().strftime("%Y-%m-%d %I:%M %p"),
            "chunks": chunks,
            "full_text_preview": full_text[:1200]
        }

        self.documents[doc_id] = doc_record
        self._save_store()
        logger.info(f"✔ Successfully indexed document '{filename}' [ID: {doc_id}, Category: {category}, Chunks: {len(chunks)}]")

        return {
            "status": "success",
            "id": doc_id,
            "filename": p.name,
            "category": category,
            "total_pages": len(pages),
            "total_chunks": len(chunks),
            "total_words": len(full_text.split()),
            "message": f"Indexed '{p.name}' ({category.replace('_', ' ').title()}) into Knowledge Vault."
        }

    # --- Retrieval & RAG ---

    def _rank_chunks(self, query: str, doc_ids: Optional[List[str]] = None, top_k: int = 4) -> List[Dict[str, Any]]:
        """Retrieves and ranks the most relevant chunks matching the user's query."""
        query_terms = re.findall(r"\w+", query.lower())
        if not query_terms:
            return []

        stopwords = {"the", "a", "an", "is", "are", "and", "or", "in", "on", "at", "to", "for", "of", "with", "this", "that", "what", "how", "why"}
        clean_terms = [t for t in query_terms if t not in stopwords and len(t) > 2]
        if not clean_terms:
            clean_terms = query_terms

        scored_chunks = []
        target_docs = [self.documents[did] for did in (doc_ids or self.documents.keys()) if did in self.documents]

        for doc in target_docs:
            for chunk in doc.get("chunks", []):
                text_lower = chunk.get("text", "").lower()
                score = 0.0

                for term in clean_terms:
                    count = text_lower.count(term)
                    if count > 0:
                        score += 1.0 + math.log(1.0 + count)

                if len(clean_terms) > 1 and " ".join(clean_terms) in text_lower:
                    score += 5.0

                if any(t in doc.get("filename", "").lower() for t in clean_terms):
                    score += 2.0

                if score > 0:
                    scored_chunks.append({
                        "score": score,
                        "doc_id": doc.get("id"),
                        "filename": doc.get("filename"),
                        "category": doc.get("category"),
                        "page": chunk.get("page"),
                        "text": chunk.get("text")
                    })

        scored_chunks.sort(key=lambda x: x["score"], reverse=True)
        return scored_chunks[:top_k]

    def answer_query(
        self,
        query: str,
        doc_id: Optional[str] = None,
        llm_client: Optional[Any] = None
    ) -> Dict[str, Any]:
        """Generates an augmented answer to a question across local documents using NVIDIA NIM."""
        if not self.documents:
            return {
                "status": "empty",
                "answer": "There are currently no documents indexed in the Knowledge Vault. Upload or specify a local PDF, code file, lecture note, or resume first!",
                "citations": []
            }

        target_ids = [doc_id] if (doc_id and doc_id in self.documents) else list(self.documents.keys())

        context_blocks = []
        citations = []

        if len(target_ids) == 1:
            doc = self.documents[target_ids[0]]
            if doc.get("total_words", 0) <= 3500:
                all_text = "\n\n".join(c["text"] for c in doc.get("chunks", []))
                context_blocks.append(f"--- Document: {doc['filename']} ({doc['category']}) ---\n{all_text}")
                citations.append({"filename": doc["filename"], "page": "All Pages"})
            else:
                top_chunks = self._rank_chunks(query, doc_ids=target_ids, top_k=5)
                for ch in top_chunks:
                    context_blocks.append(f"--- Document: {ch['filename']} [Page/Line: {ch['page']}] ---\n{ch['text']}")
                    citations.append({"filename": ch["filename"], "page": ch["page"]})
        else:
            top_chunks = self._rank_chunks(query, doc_ids=target_ids, top_k=5)
            for ch in top_chunks:
                context_blocks.append(f"--- Document: {ch['filename']} [Page/Line: {ch['page']}] ---\n{ch['text']}")
                citations.append({"filename": ch["filename"], "page": ch["page"]})

        if not context_blocks:
            latest_doc = list(self.documents.values())[-1]
            first_chunks = latest_doc.get("chunks", [])[:3]
            for ch in first_chunks:
                context_blocks.append(f"--- Document: {latest_doc['filename']} [Page: {ch['page']}] ---\n{ch['text']}")
                citations.append({"filename": latest_doc["filename"], "page": ch["page"]})

        context_text = "\n\n".join(context_blocks)

        system_prompt = (
            "You are IGIRS AI Document Intelligence Engine. You answer the user's questions about their local "
            "PDFs, lecture notes, resumes, and code files with extreme precision, citing sources accurately.\n"
            "Rules:\n"
            "1. Answer directly based on the provided document excerpts.\n"
            "2. If the user asks about a resume, highlight key experience, technologies, and achievements.\n"
            "3. If the user asks about lecture notes, explain core concepts, definitions, and equations clearly.\n"
            "4. If the user asks about code, explain logic, functions, potential bugs, or performance implications.\n"
            "5. Cite the document name and page/line number when giving facts."
        )

        user_prompt = (
            f"Here are the relevant excerpts from the user's local documents:\n\n"
            f"{context_text}\n\n"
            f"User Question: '{query}'\n\n"
            f"Provide a clear, detailed, and structured response addressing the user's query."
        )

        if not llm_client:
            from llm.nvidia_client import NvidiaLLMClient
            llm_client = NvidiaLLMClient()

        try:
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ]
            resp = llm_client.chat_completion(messages=messages, max_tokens=600)
            answer = resp.get("choices", [])[0].get("message", {}).get("content", "").strip()
            return {
                "status": "success",
                "query": query,
                "answer": answer,
                "citations": citations,
                "doc_count": len(target_ids)
            }
        except Exception as e:
            logger.error(f"RAG query execution error: {e}")
            return {
                "status": "error",
                "message": str(e),
                "answer": f"Error querying documents: {e}",
                "citations": []
            }

    def summarize_document(
        self,
        doc_id: Optional[str] = None,
        focus: Optional[str] = None,
        llm_client: Optional[Any] = None
    ) -> Dict[str, Any]:
        """Generates a comprehensive executive summary of an indexed document."""
        if not self.documents:
            return {"status": "empty", "summary": "No documents currently in Knowledge Vault to summarize."}

        target_id = doc_id if (doc_id and doc_id in self.documents) else list(self.documents.keys())[-1]
        doc = self.documents[target_id]

        chunks = doc.get("chunks", [])
        sample_text = "\n\n".join(c["text"] for c in chunks[:8])

        category = doc.get("category", "document")
        focus_instruction = f" Focus especially on: {focus}." if focus else ""

        system_prompt = (
            f"You are IGIRS AI Document Analyst. Summarize this {category.replace('_', ' ')} with crystal clarity.{focus_instruction}\n"
            "Format your response in GitHub-flavored markdown with:\n"
            "- **Executive Overview** (2-3 sentences)\n"
            "- **Key Topics / Highlights** (bullet points)\n"
            "- **Actionable Insights / Takeaways**"
        )

        user_prompt = (
            f"Document Title: {doc['filename']}\n"
            f"Document Category: {category}\n"
            f"Total Pages/Sections: {doc['total_pages']}\n\n"
            f"Document Content Excerpts:\n{sample_text}\n\n"
            f"Please generate the structured summary."
        )

        if not llm_client:
            from llm.nvidia_client import NvidiaLLMClient
            llm_client = NvidiaLLMClient()

        try:
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ]
            resp = llm_client.chat_completion(messages=messages, max_tokens=500)
            summary = resp.get("choices", [])[0].get("message", {}).get("content", "").strip()
            return {
                "status": "success",
                "doc_id": target_id,
                "filename": doc["filename"],
                "category": category,
                "summary": summary
            }
        except Exception as e:
            logger.error(f"Document summary error: {e}")
            return {"status": "error", "summary": f"Failed to summarize document: {e}"}

    def list_documents(self) -> List[Dict[str, Any]]:
        """Returns metadata list of all indexed documents."""
        out = []
        for did, d in self.documents.items():
            out.append({
                "id": did,
                "filename": d.get("filename"),
                "category": d.get("category"),
                "extension": d.get("extension"),
                "total_pages": d.get("total_pages", 1),
                "total_chunks": d.get("total_chunks", 0),
                "total_words": d.get("total_words", 0),
                "created_at": d.get("created_at", "")
            })
        return out

    def delete_document(self, doc_id: str) -> bool:
        """Removes a document from the vault."""
        if doc_id in self.documents:
            doc = self.documents.pop(doc_id)
            try:
                stored = doc.get("stored_path")
                if stored and Path(stored).exists():
                    Path(stored).unlink(missing_ok=True)
            except Exception:
                pass
            self._save_store()
            return True
        return False

    def clear_all(self) -> bool:
        """Clears all indexed documents."""
        self.documents = {}
        self._save_store()
        return True
